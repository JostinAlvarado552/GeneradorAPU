from django.shortcuts import render, redirect
from django.contrib import messages
from django.http import HttpResponse
from django.conf import settings
from docx.opc.oxml import qn, parse_xml
from docx.oxml.ns import nsdecls
from openpyxl.reader.excel import load_workbook
from .forms import IDForm
from .models import Rubros, Definicion, Medicion_Pago
from docx import Document
import xlwings as xw
from django.shortcuts import render
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet

from django.shortcuts import render
from django.contrib import messages

from django.shortcuts import render
from .models import Rubros, Definicion, Especificacion, Medicion_Pago

from django.shortcuts import render
from .models import Rubros, Definicion, Especificacion, Medicion_Pago

def buscar_rubros(request):
    if request.method == 'POST':
        ids = request.POST.get('ids', '').split(',')
        rubros_data = {}
        for id in ids:
            try:
                rubro = Rubros.objects.get(id=int(id))
                definiciones = list(Definicion.objects.filter(rubro=rubro).values('contenido'))
                especificaciones = list(Especificacion.objects.filter(rubro=rubro).values('contenido'))
                mediciones_pagos = list(Medicion_Pago.objects.filter(rubro=rubro).values('medicion', 'metodoPago'))
                detalles = process_id(id)  # Assuming you have a function to process rubro details
                rubros_data[id] = {
                    'rubro': rubro.categoria,  # Usar los campos relevantes del rubro
                    'definiciones': definiciones,
                    'especificaciones': especificaciones,
                    'mediciones_pagos': mediciones_pagos,
                    'detalles': detalles,
                }
            except Rubros.DoesNotExist:
                messages.error(request, f"No se encontró el rubro con ID {id}")
            except Exception as e:
                messages.error(request, f"Error al procesar el rubro con ID {id}: {str(e)}")

        context = {
            'rubros_data': rubros_data,
        }
        print(context)
        return render(request, 'resultado_rubros.html', context)
    else:
        return render(request, 'buscar_rubros.html')





def process_id(id_value):
    Detalles = {}
    try:
        wb = xw.Book('automatizacion_APUS/documento/APU.xlsm')
        sheet = wb.sheets['ANALISIS']
        sheet.range('C5').value = id_value
        wb.app.calculate()

        for i in range(18, 38):
            nombre = sheet.range(f'D{i}').value
            cantidad = sheet.range(f'F{i}').value
            if nombre != 0.0:
                Detalles[sheet.range('D16').value] = {'nombre': nombre, 'cantidad': cantidad}

        for i in range(41, 61):
            nombre = sheet.range(f'D{i}').value
            cantidad = sheet.range(f'F{i}').value
            if nombre != 0.0:
                Detalles[sheet.range('D39').value] = {'nombre': nombre, 'cantidad': cantidad}

        for i in range(64, 84):
            nombre = sheet.range(f'D{i}').value
            cantidad = sheet.range(f'H{i}').value
            if nombre != 0.0:
                Detalles[sheet.range('D62').value] = {'nombre': nombre, 'cantidad': cantidad}

        for i in range(87, 92):
            nombre = sheet.range(f'D{i}').value
            cantidad = sheet.range(f'H{i}').value
            if nombre != 0.0:
                Detalles[sheet.range('D85').value] = {'nombre': nombre, 'cantidad': cantidad}

        wb.save()
        wb.close()

    except Exception as e:
        return {'error': str(e)}

    return Detalles


def generar_documento_word(rubros, definiciones, medicion_pagos, detalles_dict):
    # Crear un nuevo documento
    doc = Document()
    doc.add_heading('Resultados de Búsqueda de Rubros', 0)

    for rubro in rubros:
        # Añadir título de rubro
        doc.add_heading(f'Rubro ID: {rubro.id}', level=1)

        # Crear la tabla con la primera fila de 4 columnas y las siguientes de 2 columnas
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'

        # Primera fila: 4 columnas
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'ID'
        hdr_cells[1].text = str(rubro.id)
        hdr_cells[2].text = 'UNIDAD'
        hdr_cells[3].text = rubro.especificaciones

        # Segunda fila: Nombre del rubro (2 columnas)
        table.add_row()
        row_cells = table.add_row().cells
        row_cells[0].text = 'Nombre del rubro:'
        row_cells[1].merge(row_cells[3])  # Merge cells from column 2 to 4
        row_cells[1].text = rubro.concepto

        # Tercera fila: Definición (2 columnas)
        row_cells = table.add_row().cells
        row_cells[0].text = 'Definición:'
        definiciones_text = '\n'.join([definicion.contenido for definicion in definiciones.get(rubro.id, ['No hay definiciones disponibles.'])])
        row_cells[1].merge(row_cells[3])
        row_cells[1].text = definiciones_text

        # Cuarta fila: Especificaciones (2 columnas)
        row_cells = table.add_row().cells
        row_cells[0].text = 'Especificaciones:'
        especificaciones_text = "default"
        row_cells[1].merge(row_cells[3])  # Merge cells from column 2 to 4
        row_cells[1].text = especificaciones_text

        # Quinta fila: Detalles (2 columnas)
        row_cells = table.add_row().cells
        row_cells[0].text = 'Detalles:'
        row_cells[1].merge(row_cells[3])  # Merge cells from column 2 to 4
        row_cells[1].text = detalles_dict.get(rubro.id, 'No hay detalles disponibles.')

        # Establecer el color de fondo de las celdas
        for row in table.rows:
            row.cells[0]._element.get_or_add_tcPr().append(
                parse_xml(r'<w:shd {} w:fill="D9D9D9"/>'.format(nsdecls('w'))))
            for cell in row.cells[1:]:
                cell._element.get_or_add_tcPr().append(
                    parse_xml(r'<w:shd {} w:fill="FFFFFF"/>'.format(nsdecls('w'))))

        doc.add_paragraph('')

        # Añadir mediciones y métodos de pago como una tabla
        doc.add_heading('Mediciones y Métodos de Pago', level=2)
        if medicion_pagos.get(rubro.id):
            table = doc.add_table(rows=1, cols=2)
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Medición'
            hdr_cells[1].text = 'Método de Pago'
            for mp in medicion_pagos[rubro.id]:
                row_cells = table.add_row().cells
                row_cells[0].text = mp.medicion
                row_cells[1].text = mp.metodoPago
        else:
            doc.add_paragraph('No hay mediciones y métodos de pago disponibles.')

    # Crear una respuesta con el documento
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename=resultados_rubros.docx'
    doc.save(response)
    return response
def generar_documento_pdf(rubros, definiciones, medicion_pagos, detalles_dict):
    # Crear un nuevo documento PDF
    pdf_response = HttpResponse(content_type='application/pdf')
    pdf_response['Content-Disposition'] = 'attachment; filename=resultados_rubros.pdf'

    doc = SimpleDocTemplate(pdf_response, pagesize=letter)
    styles = getSampleStyleSheet()

    elements = []

    # Contenido del PDF
    elements.append(Paragraph('Resultados de Búsqueda de Rubros', styles['Title']))

    for i, rubro in enumerate(rubros):
        if i > 0:
            elements.append(Spacer(1, 12))  # Espacio antes de cada nuevo rubro

        elements.append(Paragraph(f'Rubro ID: {rubro.id}', styles['Heading1']))
        elements.append(Paragraph(f'Concepto: {rubro.concepto}', styles['Normal']))
        elements.append(Paragraph(f'Unidad: {rubro.especificaciones}', styles['Normal']))
        elements.append(Paragraph(f'Detalles: {detalles_dict}', styles['Normal']))

        # Definiciones
        elements.append(Paragraph('Definiciones:', styles['Heading2']))
        if definiciones[rubro.id]:
            for definicion in definiciones[rubro.id]:
                elements.append(Paragraph(definicion.contenido, styles['Bullet']))
        else:
            elements.append(Paragraph('No hay definiciones disponibles.', styles['Normal']))

        # Mediciones y Métodos de Pago
        elements.append(Paragraph('Mediciones y Métodos de Pago:', styles['Heading2']))
        if medicion_pagos[rubro.id]:
            for mp in medicion_pagos[rubro.id]:
                elements.append(Paragraph(f'Medición: {mp.medicion}', styles['Bullet']))
                elements.append(Paragraph(f'Método de Pago: {mp.metodoPago}', styles['Bullet']))
        else:
            elements.append(Paragraph('No hay mediciones y métodos de pago disponibles.', styles['Normal']))

    doc.build(elements)
    return pdf_response

